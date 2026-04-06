from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO

from docx import Document


@dataclass(slots=True)
class LoadedDocument:
    text: str
    file_type: str
    styled_ranges: list[tuple[int, int]] = field(default_factory=list)


def extract_document(filename: str, payload: bytes) -> LoadedDocument:
    suffix = filename.lower().rsplit(".", 1)[-1]
    if suffix == "docx":
        return _extract_docx(payload)
    raise ValueError(f"Unsupported file type: {suffix}")


def extract_text(filename: str, payload: bytes) -> str:
    return extract_document(filename, payload).text


def _extract_docx(payload: bytes) -> LoadedDocument:
    document = Document(BytesIO(payload))
    chunks: list[str] = []
    styled_ranges: list[tuple[int, int]] = []
    cursor = 0

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        paragraph_text = paragraph.text
        if not paragraph_text.strip():
            continue

        local_cursor = 0
        for run in paragraph.runs:
            if not run.text:
                continue
            relative_start = paragraph_text.find(run.text, local_cursor)
            if relative_start == -1:
                continue
            local_cursor = relative_start + len(run.text)
            run_start = cursor + relative_start
            run_end = run_start + len(run.text)
            if run.italic or run.underline:
                styled_ranges.append((run_start, run_end))

        chunks.append(paragraph_text)
        cursor += len(paragraph_text)
        if paragraph_index != len(document.paragraphs) - 1:
            chunks.append("\n")
            cursor += 1

    return LoadedDocument(
        text="".join(chunks),
        file_type="docx",
        styled_ranges=styled_ranges,
    )
